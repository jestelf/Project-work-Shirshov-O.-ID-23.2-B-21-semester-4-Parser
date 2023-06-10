import csv
import requests
import os
import time
from bs4 import BeautifulSoup as BS
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import pandas as pd
import matplotlib.pyplot as plt
import logging
from fuzzywuzzy import process

os.system('cls' if os.name == 'nt' else 'clear')

logging.basicConfig(filename='log.txt', level=logging.ERROR, format='%(asctime)s %(message)s')

base_url = 'https://freetp.org/page/'

headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:89.0) Gecko/20100101 Firefox/89.0'
}

path = os.path.dirname(os.path.realpath(__file__)) + '\\'

os.makedirs(path + 'images', exist_ok=True)

mode = input("Выбери режимы работы программы (1 - Парсер(собирает всю информацию с нуля), 2 - Архивист(работает уже с имеющейся информацией)): ")

while True:
    try:
        threshold = int(input("Введите значение заполненности от 0 до 100. 0 - будет полное обобщение категорий, 100 - все категории будут из сайта.   Процент: "))
        if 0 <= threshold <= 100:
            break
        else:
            print("Пороговое значение должно быть между 0 и 100. Пожалуйста, попробуйте снова.")
    except ValueError:
        print("Неверный ввод. Пожалуйста, введите число.")

def unify_categories(df, column, threshold):
    unique_values = df[column].unique()
    for value in unique_values:
        matches = process.extract(value, unique_values, limit=len(unique_values))
        similar_values = [match[0] for match in matches if match[1] >= threshold]
        df.loc[df[column].isin(similar_values), column] = value
    return df

if mode == '1':
    with open(path + 'games.csv', 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)

        writer.writerow(['Title', 'Link', 'Image', 'Description', 'Comments', 'Game Mode', 'Game Language', 'Game Genre', 'Max Players', 'Single Player'])

        num_pages = int(input("Введите количество страниц для парсинга: "))

        page_number = 1

        doc = Document()

        saved_games = 0
        errors = 0

        while True:
            r = requests.get(base_url + str(page_number), headers=headers)

            if r.status_code != 200:
                error_message = f"Не удалось получить страницу {page_number}, код состояния: {r.status_code}"
                logging.error(error_message)
                break

            html = BS(r.content, 'html.parser')

            if not html.select(".base"):
                break

            for el in html.select(".base"):
                try:
                    title = el.select('.header-h1 > a > h1')
                    link = el.select('.header-h1 > a')
                    image = el.select('.short-story .maincont div img')
                    description = el.select('.short-story .maincont div p')
                    comments = el.select('.mlink .argcoms a')

                    title_text = title[0].text if title else "Заголовок отсутствует."
                    link_href = link[0]['href'] if link else "Ссылка отсутствует."
                    image_src = image[0]['src'] if image else "Картинка отсутствует."
                    description_text = description[0].text if description else "Описание отсутствует."
                    comments_count = comments[0].text if comments else "Комментарии отсутствуют"

                    image_path = path + 'images\\' + image_src.split('/')[-1]
                    try:
                        with open(image_path, 'wb') as img_file:
                            img_file.write(requests.get('https://freetp.org' + image_src).content)
                    except Exception as e:
                        error_message = f"Ошибка скачивания картинки {image_src}, error: {e}"
                        logging.error(error_message)
                        errors += 1
                        image_path = None

                    saved_games += 1
                    time.sleep(0)
                    game_page = requests.get(link_href, headers=headers)
                    if game_page.status_code == 200:
                        game_html = BS(game_page.content, 'html.parser')
                        game_mode = game_html.select_one('p:contains("Способ Игры:")') or game_html.select_one('p:contains("Способ игры:")')
                        game_language = game_html.select_one('p:contains("Язык в Игре:")') or game_html.select_one('p:contains("Язык в игре:")')
                        game_genre = game_html.select_one('p:contains("Жанр:")')
                        max_players = game_html.select_one('p:contains("Максимальное количество игроков:")') or game_html.select_one('p:contains("Количество игроков:")')
                        single_player = game_html.select_one('p:contains("Одиночная игра:")')
                        
                        game_mode_text = game_mode.text.split(":")[1].strip() if game_mode else "Способ игры не указан"
                        game_language_text = game_language.text.split(":")[1].strip() if game_language else "Язык в игре не указан"
                        game_genre_text = game_genre.text.split(":")[1].strip() if game_genre else "Жанр не указан"
                        max_players_text = max_players.text.split(":")[1].strip() if max_players else "Максимальное количество игроков не указано"
                        single_player_text = single_player.text.split(":")[1].strip() if single_player else "Информация об одиночной игре не указана"

                        writer.writerow([title_text, link_href, image_src, description_text, comments_count, game_mode_text, game_language_text, game_genre_text, max_players_text, single_player_text])

                        doc.add_heading(title_text, level=1)
                        doc.add_paragraph(f"Ссылка: {link_href}")
                        doc.add_paragraph(f"Описание: {description_text}")
                        doc.add_paragraph(f"Комментарии: {comments_count}")
                        doc.add_paragraph(f"Режим игры: {game_mode_text}")
                        doc.add_paragraph(f"Языки: {game_language_text}")
                        doc.add_paragraph(f"Жанры: {game_genre_text}")
                        doc.add_paragraph(f"Количество игроков: {max_players_text}")
                        doc.add_paragraph(f"Одиночная игра: {single_player_text}")
                        doc.add_paragraph()
                        doc.add_picture(image_path)

                    else:
                        error_message = f"Не удалось получить страницу игры {link_href}, код статуса: {game_page.status_code}"
                        logging.error(error_message)
                except Exception as e:
                    error_message = f"Не удалось обработать игру {link_href}, error: {e}"
                    logging.error(error_message)
                    errors += 1   
            
            print(f"Прошел страницу {page_number}, успешно сохранено {saved_games - errors} игр")    
            page_number += 1
            if page_number > num_pages:
                break

        doc.save(path + 'games.docx')
        pdf_path = path + 'games.pdf'
        convert(path + 'games.docx', path + 'games.pdf')
        while not os.path.exists(pdf_path):
            time.sleep(1)
        os.system("taskkill /f /im WINWORD.EXE")

df = pd.read_csv(path + 'games.csv')

print(df.describe())

df = df.drop('Game Genre', axis=1).join(df['Game Genre'].str.split(',', expand=True).stack().reset_index(level=1, drop=True).rename('Game Genre'))

df['Game Genre'] = df['Game Genre'].astype(str)
df = unify_categories(df, 'Game Genre', threshold)

df['Game Mode'] = df['Game Mode'].astype(str)
df = unify_categories(df, 'Game Mode', threshold)

df['Game Language'] = df['Game Language'].astype(str)
df = unify_categories(df, 'Game Language', threshold)

df['Max Players'] = df['Max Players'].astype(str)
df = unify_categories(df, 'Max Players', threshold)

df['Single Player'] = df['Single Player'].astype(str)
df = unify_categories(df, 'Single Player', threshold)

visualizations = {
    'Количество игр в каждом жанре': {
        'data': df['Game Genre'].value_counts(),
        'kind': 'pie'
    },
    'Количество игр в каждом игровом режиме': {
        'data': df['Game Mode'].value_counts(),
        'kind': 'pie'
    },
    'Количество игр на каждом языке': {
        'data': df['Game Language'].value_counts(),
        'kind': 'pie'
    },
    'Количество игр для каждого максимального значения игроков': {
        'data': df['Max Players'].value_counts(),
        'kind': 'pie'
    }
}

df['Comments'] = df['Comments'].replace('Комментарии отсутствуют', 0).astype(int)
popularity_levels = {
    'Популярные': 100,
    'Интересные': 50,
    'Обычные': 0
}
categories = ['Game Genre', 'Game Mode', 'Game Language', 'Max Players', 'Single Player']
categories.append('Comments')

while True:
    os.system('cls' if os.name == 'nt' else 'clear')
    print("\n1. Визуализация")
    print("2. Сортировка")
    print("0. Выход")

    main_choice = input("Укажите выбор: ")

    if main_choice == '1':
        while True:
            try:
                os.system('cls' if os.name == 'nt' else 'clear')
                for i, viz in enumerate(visualizations.keys()):
                    print(f"{i+1}. {viz}")
                viz_choice = input("Введите номер визуализации (или нажмите Enter чтобы вернуться назад): ")
                if viz_choice.lower() == '':
                    break
                viz_choice = int(viz_choice) - 1
                viz_key = list(visualizations.keys())[viz_choice]
                viz = visualizations[viz_key]

                plt.figure(figsize=(10,6))
                viz['data'].plot(kind=viz['kind'], autopct='%1.1f%%')
                plt.title(viz_key)
                plt.ylabel('')
                plt.show()

            except Exception as e:
                error_message = f"Произошла ошибка: {e}"
                logging.error(error_message)

    elif main_choice == '2':
        while True:
            try:
                os.system('cls' if os.name == 'nt' else 'clear')
                for i, cat in enumerate(categories):
                    print(f"{i+1}. {cat}")
                cat_choice = input("Введите номер категории (или нажмите Enter чтобы вернуться назад): ")
                if cat_choice.lower() == '':
                    break
                cat_choice = int(cat_choice) - 1
                cat = categories[cat_choice]

                if cat == 'Comments':
                    for i, level in enumerate(popularity_levels.keys()):
                        print(f"{i+1}. {level}")

                    level_index = int(input("Введите номер уровня популярности: ")) - 1
                    level = list(popularity_levels.keys())[level_index]

                    if level == 'Популярные':
                        filtered_df = df[df[cat] >= popularity_levels[level]]
                    elif level == 'Интересные':
                        filtered_df = df[(df[cat] < popularity_levels['Популярные']) & (df[cat] >= popularity_levels[level])]
                    elif level == 'Обычные':
                        filtered_df = df[(df[cat] > popularity_levels['Интересные']) & (df[cat] > popularity_levels[level])]
                else:
                    values = df[cat].unique()

                    for i, value in enumerate(values):
                        print(f"{i+1}. {value}")

                    value_index = int(input("Введите номер значения: ")) - 1

                    filtered_df = df[df[cat] == values[value_index]]

                print(filtered_df)

                input("Нажмите Enter, чтобы продолжить...")
            except Exception as e:
                error_message = f"Произошла ошибка: {e}"
                logging.error(error_message)

    elif main_choice == '0':
        break

    else:
        print("Неверный выбор. Пожалуйста, попробуйте снова.")
